"""
文档解析工具（增强版）。

支持格式：
- Word (.docx) - python-docx
- PDF (.pdf) - pdfplumber / Unstructured.io
- 纯文本 (.txt, .md)
- 图片（调用视觉模型识别）
- 其他格式 - Unstructured.io

新增功能：
1. Unstructured.io 专业文档处理
2. MinIO 对象存储集成
3. 表格提取
4. 页眉页脚识别
5. 智能路由：根据文件类型选择最佳解析器
6. 格式上下文保留
"""

from __future__ import annotations

import base64
import io
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from backend.config import settings


class DocumentParseError(RuntimeError):
    """文档解析异常。"""


@dataclass
class ParsedDocument:
    """
    解析后的文档结构。
    """
    text: str
    filename: str
    file_type: str
    page_count: int = 0
    has_tables: bool = False
    tables: list[list[list[str]]] = field(default_factory=list)
    has_headers_footers: bool = False
    headers: list[str] = field(default_factory=list)
    footers: list[str] = field(default_factory=list)
    elements: list[dict[str, Any]] = field(default_factory=list)
    parser_used: str = "basic"
    metadata: dict[str, Any] = field(default_factory=dict)


# ========== MinIO 对象存储 ==========

class MinIOStorage:
    """
    MinIO 对象存储客户端。

    设计原则：
    1. 懒加载连接
    2. 降级方案：MinIO不可用时使用本地文件系统
    3. 统一的上传/下载接口
    """

    def __init__(
        self,
        endpoint: Optional[str] = None,
        access_key: Optional[str] = None,
        secret_key: Optional[str] = None,
        bucket: Optional[str] = None,
        secure: Optional[bool] = None,
    ) -> None:
        self.endpoint = endpoint or settings.minio_endpoint
        self.access_key = access_key or settings.minio_access_key
        self.secret_key = (
            secret_key
            or (settings.minio_secret_key.get_secret_value() if settings.minio_secret_key else None)
        )
        self.bucket = bucket or settings.minio_bucket
        self.secure = secure if secure is not None else settings.minio_secure

        self._client = None
        self._connected = False
        self._use_local_fallback = False

        # 本地fallback目录
        self._local_fallback_dir = Path(tempfile.gettempdir()) / "lvshe_minio_fallback"
        self._local_fallback_dir.mkdir(parents=True, exist_ok=True)

    def connect(self) -> bool:
        """
        连接MinIO。

        Returns:
            True 表示连接成功，False 表示失败（使用本地fallback）
        """
        if self._connected:
            return True

        if not settings.minio_enabled:
            self._use_local_fallback = True
            return False

        if not self.access_key or not self.secret_key:
            print("[DocumentParser] MinIO未配置密钥，使用本地存储")
            self._use_local_fallback = True
            return False

        try:
            from minio import Minio

            self._client = Minio(
                self.endpoint,
                access_key=self.access_key,
                secret_key=self.secret_key,
                secure=self.secure,
            )

            # 确保bucket存在
            if not self._client.bucket_exists(self.bucket):
                self._client.make_bucket(self.bucket)

            self._connected = True
            print("[DocumentParser] MinIO连接成功")
            return True

        except Exception as exc:
            print(f"[DocumentParser] MinIO连接失败，使用本地存储: {exc}")
            self._use_local_fallback = True
            return False

    def upload_file(self, file_path: str | Path, object_name: str) -> str:
        """
        上传文件到MinIO。

        Returns:
            对象名称
        """
        if self._use_local_fallback or not self.connect():
            # 本地fallback
            local_path = self._local_fallback_dir / object_name
            local_path.parent.mkdir(parents=True, exist_ok=True)
            import shutil
            shutil.copy2(file_path, local_path)
            return object_name

        try:
            self._client.fput_object(
                self.bucket,
                object_name,
                str(file_path),
            )
            return object_name
        except Exception as exc:
            print(f"[DocumentParser] MinIO上传失败，使用本地存储: {exc}")
            # 降级到本地
            local_path = self._local_fallback_dir / object_name
            local_path.parent.mkdir(parents=True, exist_ok=True)
            import shutil
            shutil.copy2(file_path, local_path)
            return object_name

    def download_file(self, object_name: str, local_path: str | Path) -> str:
        """
        从MinIO下载文件。

        Returns:
            本地文件路径
        """
        if self._use_local_fallback or not self.connect():
            local_file = self._local_fallback_dir / object_name
            if local_file.exists():
                import shutil
                shutil.copy2(local_file, local_path)
                return str(local_path)
            raise DocumentParseError(f"文件不存在: {object_name}")

        try:
            self._client.fget_object(
                self.bucket,
                object_name,
                str(local_path),
            )
            return str(local_path)
        except Exception as exc:
            raise DocumentParseError(f"MinIO下载失败: {exc}") from exc

    def file_exists(self, object_name: str) -> bool:
        """检查文件是否存在。"""
        if self._use_local_fallback:
            return (self._local_fallback_dir / object_name).exists()

        try:
            self._client.stat_object(self.bucket, object_name)
            return True
        except Exception:
            return False


# ========== Unstructured.io 解析器 ==========

class UnstructuredParser:
    """
    Unstructured.io 专业文档解析器。

    支持：
    - PDF、DOCX、PPTX、XLSX等多种格式
    - 表格提取
    - 页眉页脚识别
    - 元素级解析

    降级方案：Unstructured不可用时回退到基础解析器
    """

    def __init__(
        self,
        api_url: Optional[str] = None,
        api_key: Optional[str] = None,
    ) -> None:
        self.api_url = api_url or settings.unstructured_api_url
        self.api_key = (
            api_key
            or (settings.unstructured_api_key.get_secret_value() if settings.unstructured_api_key else None)
        )

        self._local_available = False
        self._api_available = bool(self.api_url and self.api_key)

        # 检查本地库是否可用
        self._check_local_available()

    def _check_local_available(self) -> None:
        """检查本地unstructured库是否可用。"""
        try:
            import unstructured  # noqa: F401
            self._local_available = True
        except ImportError:
            self._local_available = False

    @property
    def is_available(self) -> bool:
        return self._local_available or self._api_available

    def parse_file(self, file_path: str | Path) -> ParsedDocument:
        """
        解析文件。

        优先使用本地库，其次使用API，都不可用时抛出异常。
        """
        path = Path(file_path)

        if self._local_available:
            return self._parse_local(path)
        elif self._api_available:
            return self._parse_api(path)
        else:
            raise DocumentParseError("Unstructured不可用：本地库和API都未配置")

    def _parse_local(self, path: Path) -> ParsedDocument:
        """使用本地unstructured库解析。"""
        try:
            from unstructured.partition.auto import partition

            elements = partition(filename=str(path))

            # 提取文本
            text_parts = []
            tables = []
            headers = []
            footers = []
            elements_data = []

            for element in elements:
                element_dict = element.to_dict() if hasattr(element, "to_dict") else {"text": str(element)}
                elements_data.append(element_dict)

                text = str(element)
                text_parts.append(text)

                # 分类元素
                category = element_dict.get("type", "")
                if category == "Table":
                    # 尝试提取表格数据
                    tables.append(self._extract_table_data(element))
                elif category == "Header":
                    headers.append(text)
                elif category == "Footer":
                    footers.append(text)

            return ParsedDocument(
                text="\n\n".join(text_parts),
                filename=path.name,
                file_type=path.suffix.lower(),
                page_count=self._count_pages(elements),
                has_tables=len(tables) > 0,
                tables=tables,
                has_headers_footers=len(headers) > 0 or len(footers) > 0,
                headers=headers,
                footers=footers,
                elements=elements_data,
                parser_used="unstructured_local",
            )

        except ImportError as exc:
            raise DocumentParseError(f"Unstructured本地解析失败: {exc}") from exc
        except Exception as exc:
            raise DocumentParseError(f"Unstructured解析失败: {exc}") from exc

    def _parse_api(self, path: Path) -> ParsedDocument:
        """使用Unstructured API解析。"""
        try:
            import requests

            headers = {}
            if self.api_key:
                headers["unstructured-api-key"] = self.api_key

            with open(path, "rb") as f:
                files = {"files": (path.name, f)}
                response = requests.post(
                    self.api_url,
                    headers=headers,
                    files=files,
                    timeout=60,
                )

            if response.status_code != 200:
                raise DocumentParseError(
                    f"Unstructured API调用失败: {response.status_code} - {response.text}"
                )

            elements = response.json()

            # 解析结果
            text_parts = []
            tables = []
            headers = []
            footers = []

            for element in elements:
                text = element.get("text", "")
                text_parts.append(text)

                category = element.get("type", "")
                if category == "Table":
                    tables.append([[text]])  # 简化处理
                elif category == "Header":
                    headers.append(text)
                elif category == "Footer":
                    footers.append(text)

            return ParsedDocument(
                text="\n\n".join(text_parts),
                filename=path.name,
                file_type=path.suffix.lower(),
                page_count=0,
                has_tables=len(tables) > 0,
                tables=tables,
                has_headers_footers=len(headers) > 0 or len(footers) > 0,
                headers=headers,
                footers=footers,
                elements=elements,
                parser_used="unstructured_api",
            )

        except Exception as exc:
            raise DocumentParseError(f"Unstructured API解析失败: {exc}") from exc

    def _extract_table_data(self, element) -> list[list[str]]:
        """从表格元素中提取数据。"""
        try:
            if hasattr(element, "metadata") and hasattr(element.metadata, "text_as_html"):
                # 从HTML中解析表格
                html = element.metadata.text_as_html
                return self._parse_html_table(html)
        except Exception:
            pass

        # 降级：返回纯文本
        return [[str(element)]]

    def _parse_html_table(self, html: str) -> list[list[str]]:
        """简单的HTML表格解析。"""
        try:
            from html.parser import HTMLParser

            rows = []
            current_row = []
            current_cell = ""
            in_cell = False

            class TableParser(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.rows = []
                    self.current_row = []
                    self.current_cell = ""
                    self.in_cell = False

                def handle_starttag(self, tag, attrs):
                    if tag == "tr":
                        self.current_row = []
                    elif tag in ("td", "th"):
                        self.in_cell = True
                        self.current_cell = ""

                def handle_endtag(self, tag):
                    if tag in ("td", "th"):
                        self.in_cell = False
                        self.current_row.append(self.current_cell.strip())
                    elif tag == "tr":
                        if self.current_row:
                            self.rows.append(self.current_row)

                def handle_data(self, data):
                    if self.in_cell:
                        self.current_cell += data

            parser = TableParser()
            parser.feed(html)
            return parser.rows

        except Exception:
            return []

    def _count_pages(self, elements) -> int:
        """统计页数。"""
        pages = set()
        for element in elements:
            try:
                if hasattr(element, "metadata") and hasattr(element.metadata, "page_number"):
                    pages.add(element.metadata.page_number)
            except Exception:
                pass
        return len(pages) if pages else 0


# ========== 增强版文档解析器 ==========

class DocumentParser:
    """
    统一文档解析器（增强版）。

    智能路由：
    - .docx: 优先 python-docx（快），可选 Unstructured（更完整）
    - .pdf: 优先 Unstructured（专业），降级 pdfplumber
    - .txt/.md: 直接读取
    - 图片: 视觉模型识别
    - 其他: Unstructured（如果可用）

    新增功能：
    - MinIO 对象存储
    - 表格提取
    - 页眉页脚识别
    - 结构化解析结果
    """

    SUPPORTED_EXTENSIONS = {
        ".docx", ".pdf", ".txt", ".md",
        ".png", ".jpg", ".jpeg", ".bmp", ".webp",
        ".pptx", ".xlsx", ".csv",
    }

    def __init__(self, llm_gateway=None) -> None:
        self.llm_gateway = llm_gateway

        # 子组件（懒加载）
        self._unstructured_parser: Optional[UnstructuredParser] = None
        self._minio_storage: Optional[MinIOStorage] = None

    @property
    def unstructured_parser(self) -> Optional[UnstructuredParser]:
        if self._unstructured_parser is None and settings.unstructured_enabled:
            try:
                self._unstructured_parser = UnstructuredParser()
            except Exception:
                self._unstructured_parser = None
        return self._unstructured_parser

    @property
    def minio_storage(self) -> Optional[MinIOStorage]:
        if self._minio_storage is None and settings.minio_enabled:
            try:
                self._minio_storage = MinIOStorage()
            except Exception:
                self._minio_storage = None
        return self._minio_storage

    def parse_file(self, file_path: str | Path) -> str:
        """
        解析文件，返回纯文本内容。

        保持向后兼容的接口。
        """
        result = self.parse_file_structured(file_path)
        return result.text

    def parse_file_structured(self, file_path: str | Path) -> ParsedDocument:
        """
        解析文件，返回结构化结果。

        Args:
            file_path: 文件路径

        Returns:
            ParsedDocument 结构化解析结果

        Raises:
            DocumentParseError: 解析失败
        """
        path = Path(file_path)
        if not path.exists():
            raise DocumentParseError(f"文件不存在：{file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise DocumentParseError(f"不支持的文件格式：{ext}")

        try:
            # 智能路由：选择最佳解析器
            if ext == ".docx":
                return self._parse_docx_smart(path)
            elif ext == ".pdf":
                return self._parse_pdf_smart(path)
            elif ext in (".txt", ".md"):
                return self._parse_text(path)
            elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".webp"):
                return self._parse_image(path)
            elif ext in (".pptx", ".xlsx", ".csv"):
                return self._parse_with_unstructured(path)
            else:
                raise DocumentParseError(f"不支持的格式：{ext}")
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError(f"解析文件失败：{exc}") from exc

    def parse_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        从字节流解析文件。

        保持向后兼容的接口。
        """
        result = self.parse_bytes_structured(file_bytes, filename)
        return result.text

    def parse_bytes_structured(self, file_bytes: bytes, filename: str) -> ParsedDocument:
        """
        从字节流解析文件，返回结构化结果。
        """
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise DocumentParseError(f"不支持的文件格式：{ext}")

        # 写入临时文件
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            return self.parse_file_structured(tmp_path)
        finally:
            # 清理临时文件
            try:
                Path(tmp_path).unlink(missing_ok=True)
            except Exception:
                pass

    # ========== 智能解析路由 ==========

    def _parse_docx_smart(self, path: Path) -> ParsedDocument:
        """
        智能解析docx。

        策略：先用python-docx快速解析，如果需要表格等结构化信息再用Unstructured。
        """
        # 优先使用Unstructured（如果可用且需要结构化信息）
        if self.unstructured_parser and self.unstructured_parser.is_available:
            try:
                result = self.unstructured_parser.parse_file(path)
                if result.has_tables or result.has_headers_footers:
                    return result
            except Exception:
                # Unstructured失败，降级到python-docx
                pass

        # 使用python-docx
        try:
            from docx import Document

            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # 尝试提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

            return ParsedDocument(
                text="\n".join(paragraphs),
                filename=path.name,
                file_type=".docx",
                has_tables=len(tables) > 0,
                tables=tables,
                parser_used="python-docx",
            )
        except ImportError as exc:
            raise DocumentParseError(
                "缺少 python-docx 依赖，请安装：uv add python-docx"
            ) from exc
        except Exception as exc:
            raise DocumentParseError(f"Word 解析失败：{exc}") from exc

    def _parse_pdf_smart(self, path: Path) -> ParsedDocument:
        """
        智能解析PDF。

        策略：优先Unstructured（专业），降级pdfplumber。
        """
        # 优先使用Unstructured
        if self.unstructured_parser and self.unstructured_parser.is_available:
            try:
                return self.unstructured_parser.parse_file(path)
            except Exception:
                pass

        # 降级到pdfplumber
        try:
            import pdfplumber

            text_parts = []
            tables = []
            page_count = 0

            with pdfplumber.open(str(path)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                    # 提取表格
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)

            return ParsedDocument(
                text="\n\n".join(text_parts),
                filename=path.name,
                file_type=".pdf",
                page_count=page_count,
                has_tables=len(tables) > 0,
                tables=tables,
                parser_used="pdfplumber",
            )
        except ImportError as exc:
            raise DocumentParseError(
                "缺少 pdfplumber 依赖，请安装：uv add pdfplumber"
            ) from exc
        except Exception as exc:
            raise DocumentParseError(f"PDF 解析失败：{exc}") from exc

    def _parse_with_unstructured(self, path: Path) -> ParsedDocument:
        """使用Unstructured解析（pptx/xlsx/csv等）。"""
        if self.unstructured_parser and self.unstructured_parser.is_available:
            return self.unstructured_parser.parse_file(path)
        else:
            raise DocumentParseError(
                f"格式 {path.suffix} 需要 Unstructured.io 支持，请安装或配置"
            )

    # ========== 基础解析（保持向后兼容） ==========

    def _parse_text(self, path: Path) -> ParsedDocument:
        """解析纯文本文件。"""
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # 尝试 GBK
            try:
                text = path.read_text(encoding="gbk")
            except Exception as exc:
                raise DocumentParseError(f"文本文件编码无法识别：{exc}") from exc

        return ParsedDocument(
            text=text,
            filename=path.name,
            file_type=path.suffix.lower(),
            parser_used="plain_text",
        )

    def _parse_image(self, path: Path) -> ParsedDocument:
        """
        使用视觉模型识别图片中的文字。
        """
        if not self.llm_gateway:
            raise DocumentParseError(
                "图片识别需要配置 LLM 网关（视觉模型）"
            )

        try:
            # 读取图片并转 base64
            image_bytes = path.read_bytes()
            base64_image = base64.b64encode(image_bytes).decode("utf-8")

            # 调用视觉模型
            text = self._call_vision_model(base64_image, path.suffix)

            return ParsedDocument(
                text=text,
                filename=path.name,
                file_type=path.suffix.lower(),
                parser_used="vision_model",
            )
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError(f"图片识别失败：{exc}") from exc

    def _call_vision_model(self, base64_image: str, ext: str) -> str:
        """
        调用视觉模型识别图片文字。

        预留接口，具体实现根据使用的视觉模型调整。
        """
        # 预留实现：调用视觉模型进行 OCR
        # 实际项目中需要根据使用的视觉模型 API 来实现

        return (
            "【图片识别功能】\n"
            "图片识别需要配置视觉模型 API（如智谱 GLM-4V Flash）。\n"
            "请在 .env 中配置 GLM-4V 的 API Key，并在 LLM 网关中添加视觉模型支持。\n"
            "当前仅支持文本类合同（Word/PDF/TXT）的直接解析。"
        )

    # ========== MinIO 相关功能 ==========

    def upload_to_storage(self, file_path: str | Path, object_name: Optional[str] = None) -> str:
        """
        上传文件到对象存储。

        Args:
            file_path: 本地文件路径
            object_name: 对象名称（可选，默认使用文件名）

        Returns:
            对象名称
        """
        path = Path(file_path)
        if object_name is None:
            import uuid
            object_name = f"documents/{uuid.uuid4().hex[:12]}/{path.name}"

        if self.minio_storage:
            return self.minio_storage.upload_file(path, object_name)
        else:
            # 没有MinIO时返回本地路径
            return str(path)

    def download_from_storage(self, object_name: str, local_path: Optional[str | Path] = None) -> str:
        """
        从对象存储下载文件。

        Args:
            object_name: 对象名称
            local_path: 本地保存路径（可选）

        Returns:
            本地文件路径
        """
        if local_path is None:
            local_path = tempfile.mktemp(suffix=Path(object_name).suffix)

        if self.minio_storage:
            return self.minio_storage.download_file(object_name, local_path)
        else:
            # 没有MinIO时假设object_name就是本地路径
            return object_name

    def get_parser_info(self) -> dict[str, Any]:
        """获取解析器信息。"""
        return {
            "supported_extensions": list(self.SUPPORTED_EXTENSIONS),
            "unstructured_available": (
                self.unstructured_parser.is_available
                if self.unstructured_parser
                else False
            ),
            "minio_enabled": settings.minio_enabled,
            "minio_connected": (
                self.minio_storage._connected
                if self.minio_storage
                else False
            ),
        }


# 全局单例
document_parser = DocumentParser()


if __name__ == "__main__":
    # 简单测试
    parser = DocumentParser()
    print("文档解析器信息:")
    print(parser.get_parser_info())
